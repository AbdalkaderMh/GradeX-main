from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def generate_arabic_docs():
    doc = Document()

    # Helper function to add right-aligned Arabic text
    def add_arabic_heading(text, level):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_arabic_paragraph(text):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Title
    add_arabic_heading('توثيق منصة GradeX لإدارة النتائج', 0)

    # Overview
    add_arabic_heading('1. نظرة عامة على المشروع', 1)
    add_arabic_paragraph(
        'GradeX هي منصة تعليمية متكاملة مصممة لتبسيط عملية إدارة نتائج الطلاب، '
        'وتعزيز التواصل بين الإدارة والمعلمين والطلاب. توفر المنصة واجهة عصرية وسهلة الاستخدام '
        'تدعم اللغة العربية بشكل كامل.'
    )

    # Key Features
    add_arabic_heading('2. المميزات الرئيسية', 1)
    add_arabic_paragraph('• نظام إدارة مستخدمين متعدد الأدوار (مسؤول، معلم، طالب).')
    add_arabic_paragraph('• رفع النتائج عبر ملفات Excel مع ميزة المعاينة قبل النشر.')
    add_arabic_paragraph('• استخراج تقارير وشهادات ملونة للطلاب.')
    add_arabic_paragraph('• نظام إشعارات فوري لإبلاغ الطلاب بالنتائج الجديدة.')
    add_arabic_paragraph('• إحصائيات متقدمة وتحليل للأداء الأكاديمي.')
    add_arabic_paragraph('• دعم كامل لرفع وتعديل الصور الشخصية وشعار المنصة مع ميزة القص.')

    # Installation
    add_arabic_heading('3. تعليمات التثبيت', 1)
    add_arabic_heading('المتطلبات:', 2)
    add_arabic_paragraph('Node.js, MongoDB')
    add_arabic_heading('خطوات التشغيل:', 2)
    add_arabic_paragraph('1. قم بتحميل المشروع وفك الضغط.')
    add_arabic_paragraph('2. في مجلد backend: قم بتنفيذ npm install ثم npm run dev.')
    add_arabic_paragraph('3. في مجلد frontend: قم بتنفيذ npm install ثم npm run dev.')

    # User Roles
    add_arabic_heading('4. أدوار المستخدمين', 1)

    add_arabic_heading('المسؤول (Admin):', 2)
    add_arabic_paragraph('يتحكم في إعدادات المنصة، يضيف المعلمين، يعتمد الحسابات، ويطلع على الإحصائيات العامة.')

    add_arabic_heading('المعلم (Teacher):', 2)
    add_arabic_paragraph('يقوم برفع درجات المواد المسندة إليه، يتابع أداء طلابه، ويتواصل مع الإدارة.')

    add_arabic_heading('الطالب (Student):', 2)
    add_arabic_paragraph('يطلع على نتائجه، يحمل الشهادات، ويستلم الإشعارات.')

    # Technical Details
    add_arabic_heading('5. التفاصيل التقنية', 1)
    add_arabic_paragraph('• الواجهة الأمامية: React.js, Tailwind CSS.')
    add_arabic_paragraph('• الواجهة الخلفية: Node.js, Express.js.')
    add_arabic_paragraph('• قاعدة البيانات: MongoDB.')
    add_arabic_paragraph('• معالجة الصور: Sharp.')
    add_arabic_paragraph('• معالجة البيانات: XLSX.')

    doc.save('GradeX_Documentation_AR.docx')

if __name__ == '__main__':
    generate_arabic_docs()
